"""
Document parsing service.

Extracts plain text from resumes uploaded as PDF, DOCX, or TXT. PDF
parsing uses pdfplumber first (best layout fidelity for resumes, which
often use multi-column or table layouts) and falls back to PyPDF2 if
pdfplumber fails to extract anything useful (e.g. on malformed PDFs).
"""
from __future__ import annotations

import io
import re

from app.core.exceptions import InvalidFileError, TextExtractionError

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None


SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


def get_extension(filename: str) -> str:
    if "." not in filename:
        raise InvalidFileError(f"File '{filename}' has no extension.")
    ext = filename.rsplit(".", 1)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise InvalidFileError(
            f"Unsupported file type '.{ext}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )
    return ext


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = get_extension(filename)

    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    else:  # txt
        text = _extract_txt(file_bytes)

    text = clean_text(text)
    if not text.strip():
        raise TextExtractionError(
            "No readable text could be extracted from this file. "
            "It may be a scanned image without OCR, password-protected, or corrupted."
        )
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    text_chunks: list[str] = []

    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_chunks.append(page_text)
            combined = "\n".join(text_chunks)
            if combined.strip():
                return combined
        except Exception:
            # Fall through to PyPDF2 fallback below.
            text_chunks = []

    if PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text_chunks.append(page.extract_text() or "")
            combined = "\n".join(text_chunks)
            if combined.strip():
                return combined
        except Exception as exc:
            raise TextExtractionError(f"Failed to parse PDF: {exc}") from exc

    raise TextExtractionError("PDF parsing libraries are unavailable.")


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx  # python-docx, lazy import (optional dependency)
    except ImportError as exc:
        raise TextExtractionError(
            "DOCX support requires the 'python-docx' package."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of tables, which many resume templates use.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        raise TextExtractionError(f"Failed to parse DOCX: {exc}") from exc


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TextExtractionError("Could not decode text file with utf-8 or latin-1.")


def clean_text(text: str) -> str:
    """Normalize whitespace and strip control characters."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
